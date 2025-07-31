from flask import Blueprint, request, jsonify
from src.models.user import db
from src.models.transaction import Transaction, Category, Budget
from datetime import datetime, date
import pandas as pd
from io import BytesIO
import openpyxl
from docx import Document
from docx.shared import Inches
import os

finance_bp = Blueprint('finance', __name__)

# Rotas para Categorias
@finance_bp.route('/categories', methods=['GET'])
def get_categories():
    categories = Category.query.all()
    return jsonify([category.to_dict() for category in categories])

@finance_bp.route('/categories', methods=['POST'])
def create_category():
    data = request.get_json()
    
    category = Category(
        name=data['name'],
        type=data['type'],
        color=data.get('color', '#6b7280')
    )
    
    try:
        db.session.add(category)
        db.session.commit()
        return jsonify(category.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': 'Categoria já existe'}), 400

@finance_bp.route('/categories/<int:category_id>', methods=['DELETE'])
def delete_category(category_id):
    category = Category.query.get_or_404(category_id)
    
    # Verificar se há transações associadas
    if category.transactions:
        return jsonify({'error': 'Não é possível excluir categoria com transações associadas'}), 400
    
    db.session.delete(category)
    db.session.commit()
    return '', 204

# Rotas para Transações
@finance_bp.route('/transactions', methods=['GET'])
def get_transactions():
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 50, type=int)
    category_id = request.args.get('category_id', type=int)
    transaction_type = request.args.get('type')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    query = Transaction.query
    
    if category_id:
        query = query.filter(Transaction.category_id == category_id)
    if transaction_type:
        query = query.filter(Transaction.type == transaction_type)
    if start_date:
        query = query.filter(Transaction.date >= datetime.strptime(start_date, '%Y-%m-%d').date())
    if end_date:
        query = query.filter(Transaction.date <= datetime.strptime(end_date, '%Y-%m-%d').date())
    
    transactions = query.order_by(Transaction.date.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    
    return jsonify({
        'transactions': [transaction.to_dict() for transaction in transactions.items],
        'total': transactions.total,
        'pages': transactions.pages,
        'current_page': page
    })

@finance_bp.route('/transactions', methods=['POST'])
def create_transaction():
    data = request.get_json()
    
    transaction = Transaction(
        description=data['description'],
        amount=float(data['amount']),
        type=data['type'],
        category_id=data['category_id'],
        date=datetime.strptime(data['date'], '%Y-%m-%d').date() if data.get('date') else date.today()
    )
    
    db.session.add(transaction)
    db.session.commit()
    return jsonify(transaction.to_dict()), 201

@finance_bp.route('/transactions/<int:transaction_id>', methods=['PUT'])
def update_transaction(transaction_id):
    transaction = Transaction.query.get_or_404(transaction_id)
    data = request.get_json()
    
    transaction.description = data.get('description', transaction.description)
    transaction.amount = float(data.get('amount', transaction.amount))
    transaction.type = data.get('type', transaction.type)
    transaction.category_id = data.get('category_id', transaction.category_id)
    if data.get('date'):
        transaction.date = datetime.strptime(data['date'], '%Y-%m-%d').date()
    
    db.session.commit()
    return jsonify(transaction.to_dict())

@finance_bp.route('/transactions/<int:transaction_id>', methods=['DELETE'])
def delete_transaction(transaction_id):
    transaction = Transaction.query.get_or_404(transaction_id)
    db.session.delete(transaction)
    db.session.commit()
    return '', 204

# Rotas para Orçamentos
@finance_bp.route('/budgets', methods=['GET'])
def get_budgets():
    month = request.args.get('month', datetime.now().month, type=int)
    year = request.args.get('year', datetime.now().year, type=int)
    
    budgets = Budget.query.filter_by(month=month, year=year).all()
    return jsonify([budget.to_dict() for budget in budgets])

@finance_bp.route('/budgets', methods=['POST'])
def create_budget():
    data = request.get_json()
    
    budget = Budget(
        category_id=data['category_id'],
        amount=float(data['amount']),
        month=data['month'],
        year=data['year']
    )
    
    try:
        db.session.add(budget)
        db.session.commit()
        return jsonify(budget.to_dict()), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': 'Orçamento já existe para esta categoria e período'}), 400

# Rotas para Relatórios e Análises
@finance_bp.route('/dashboard', methods=['GET'])
def get_dashboard():
    month = request.args.get('month', datetime.now().month, type=int)
    year = request.args.get('year', datetime.now().year, type=int)
    
    # Transações do mês atual
    start_date = date(year, month, 1)
    if month == 12:
        end_date = date(year + 1, 1, 1)
    else:
        end_date = date(year, month + 1, 1)
    
    transactions = Transaction.query.filter(
        Transaction.date >= start_date,
        Transaction.date < end_date
    ).all()
    
    # Calcular totais
    total_income = sum(t.amount for t in transactions if t.type == 'income')
    total_expenses = sum(t.amount for t in transactions if t.type == 'expense')
    balance = total_income - total_expenses
    
    # Gastos por categoria
    expenses_by_category = {}
    for transaction in transactions:
        if transaction.type == 'expense':
            category_name = transaction.category.name
            if category_name not in expenses_by_category:
                expenses_by_category[category_name] = 0
            expenses_by_category[category_name] += transaction.amount
    
    # Orçamentos vs gastos reais
    budgets = Budget.query.filter_by(month=month, year=year).all()
    budget_analysis = []
    for budget in budgets:
        spent = sum(t.amount for t in transactions 
                   if t.type == 'expense' and t.category_id == budget.category_id)
        budget_analysis.append({
            'category': budget.category.name,
            'budgeted': budget.amount,
            'spent': spent,
            'remaining': budget.amount - spent,
            'percentage': (spent / budget.amount * 100) if budget.amount > 0 else 0
        })
    
    return jsonify({
        'total_income': total_income,
        'total_expenses': total_expenses,
        'balance': balance,
        'expenses_by_category': expenses_by_category,
        'budget_analysis': budget_analysis,
        'month': month,
        'year': year
    })

# Rota para exportar para Excel
@finance_bp.route('/export/excel', methods=['GET'])
def export_excel():
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    query = Transaction.query
    if start_date:
        query = query.filter(Transaction.date >= datetime.strptime(start_date, '%Y-%m-%d').date())
    if end_date:
        query = query.filter(Transaction.date <= datetime.strptime(end_date, '%Y-%m-%d').date())
    
    transactions = query.order_by(Transaction.date.desc()).all()
    
    # Criar DataFrame
    data = []
    for transaction in transactions:
        data.append({
            'Data': transaction.date,
            'Descrição': transaction.description,
            'Categoria': transaction.category.name,
            'Tipo': 'Receita' if transaction.type == 'income' else 'Despesa',
            'Valor': transaction.amount
        })
    
    df = pd.DataFrame(data)
    
    # Criar arquivo Excel
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Transações', index=False)
        
        # Adicionar formatação
        workbook = writer.book
        worksheet = writer.sheets['Transações']
        
        # Formatação de cabeçalho
        header_font = openpyxl.styles.Font(bold=True)
        for cell in worksheet[1]:
            cell.font = header_font
        
        # Ajustar largura das colunas
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    
    # Salvar arquivo temporário
    filename = f"transacoes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    filepath = os.path.join('/tmp', filename)
    with open(filepath, 'wb') as f:
        f.write(output.getvalue())
    
    return jsonify({'download_url': f'/api/download/{filename}'})

# Rota para gerar relatório em Word
@finance_bp.route('/export/word', methods=['GET'])
def export_word():
    month = request.args.get('month', datetime.now().month, type=int)
    year = request.args.get('year', datetime.now().year, type=int)
    
    # Obter dados do dashboard
    dashboard_data = get_dashboard().get_json()
    
    # Criar documento Word
    doc = Document()
    
    # Título
    title = doc.add_heading(f'Relatório Financeiro - {month:02d}/{year}', 0)
    
    # Resumo
    doc.add_heading('Resumo Financeiro', level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ['Total de Receitas', f"R$ {dashboard_data['total_income']:.2f}"],
        ['Total de Despesas', f"R$ {dashboard_data['total_expenses']:.2f}"],
        ['Saldo', f"R$ {dashboard_data['balance']:.2f}"],
        ['Status', 'Positivo' if dashboard_data['balance'] >= 0 else 'Negativo']
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = value
    
    # Gastos por categoria
    if dashboard_data['expenses_by_category']:
        doc.add_heading('Gastos por Categoria', level=1)
        category_table = doc.add_table(rows=len(dashboard_data['expenses_by_category']) + 1, cols=2)
        category_table.style = 'Table Grid'
        
        # Cabeçalho
        category_table.cell(0, 0).text = 'Categoria'
        category_table.cell(0, 1).text = 'Valor'
        
        for i, (category, amount) in enumerate(dashboard_data['expenses_by_category'].items(), 1):
            category_table.cell(i, 0).text = category
            category_table.cell(i, 1).text = f"R$ {amount:.2f}"
    
    # Análise de orçamentos
    if dashboard_data['budget_analysis']:
        doc.add_heading('Análise de Orçamentos', level=1)
        budget_table = doc.add_table(rows=len(dashboard_data['budget_analysis']) + 1, cols=4)
        budget_table.style = 'Table Grid'
        
        # Cabeçalho
        headers = ['Categoria', 'Orçado', 'Gasto', 'Restante']
        for i, header in enumerate(headers):
            budget_table.cell(0, i).text = header
        
        for i, budget in enumerate(dashboard_data['budget_analysis'], 1):
            budget_table.cell(i, 0).text = budget['category']
            budget_table.cell(i, 1).text = f"R$ {budget['budgeted']:.2f}"
            budget_table.cell(i, 2).text = f"R$ {budget['spent']:.2f}"
            budget_table.cell(i, 3).text = f"R$ {budget['remaining']:.2f}"
    
    # Salvar arquivo
    filename = f"relatorio_{month:02d}_{year}.docx"
    filepath = os.path.join('/tmp', filename)
    doc.save(filepath)
    
    return jsonify({'download_url': f'/api/download/{filename}'})

# Rota para download de arquivos
@finance_bp.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    from flask import send_file
    filepath = os.path.join('/tmp', filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return jsonify({'error': 'Arquivo não encontrado'}), 404

